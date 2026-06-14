"""Unit tests for the non-LLM pipeline layers. Run: .venv\\Scripts\\python -m pytest tests -q"""

import math
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from package import clean_letter, slugify
from pipeline import db
from pipeline.discovery import _annualize, _s
from pipeline.scoring import salary_gate, score_job, title_match


# ---- discovery normalization ----

def test_nan_coerced_to_empty():
    assert _s(float("nan")) == ""
    assert _s(None) == ""
    assert _s("Acme") == "Acme"


def test_annualize_intervals():
    assert _annualize(55, "hourly") == 55 * 2080
    assert _annualize(10000, "monthly") == 120000
    assert _annualize(120000, "yearly") == 120000
    assert _annualize(None, "yearly") is None
    assert _annualize(float("nan"), "hourly") is None
    assert _annualize(2000, "weekly") == 104000


def test_annualize_unknown_interval_defaults_yearly():
    assert _annualize(90000, "fortnightly") == 90000


# ---- dedup ----

def test_dedup_normalizes_company_suffixes_and_parens():
    a = db.dedup_hash("Acme Inc", "QA Engineer (Remote)")
    b = db.dedup_hash("Acme", "QA Engineer")
    assert a == b


def test_dedup_distinct_titles_differ():
    assert db.dedup_hash("Acme", "QA Engineer") != db.dedup_hash("Acme", "QA Manager")


def test_dedup_unicode_and_symbols():
    h = db.dedup_hash("Café & Co.", "Ingénieur QA — Sénior")
    assert isinstance(h, str) and len(h) == 32


def test_upsert_never_reapplies():
    """Critical invariant: re-discovering an applied job must not reset it or duplicate."""
    import sqlite3
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    conn.executescript(db.SCHEMA)
    for m in db.MIGRATIONS:
        try:
            conn.execute(m)
        except sqlite3.OperationalError:
            pass
    h = db.dedup_hash("IdemCo", "Idem QA")
    assert db.upsert_job(conn, {"title": "Idem QA", "company": "IdemCo"}) is True
    conn.execute("UPDATE jobs SET status='applied' WHERE dedup_hash=?", (h,))
    # variant phrasing of the same role must dedup to the same row
    assert db.upsert_job(conn, {"title": "Idem QA (Remote)", "company": "IdemCo Inc"}) is False
    row = conn.execute("SELECT status FROM jobs WHERE dedup_hash=?", (h,)).fetchone()
    assert row["status"] == "applied"
    assert conn.execute("SELECT COUNT(*) c FROM jobs WHERE dedup_hash=?", (h,)).fetchone()["c"] == 1


# ---- scoring ----

PROFILE = {
    "search_criteria": {"min_salary_usd": 100000},
    "personas": {
        "technical": {"target_titles": ["QA Engineer", "SDET"]},
        "analyst_pm": {"target_titles": ["Business Analyst"]},
    },
}


def _job(**kw):
    base = {"title": "QA Engineer", "company": "X", "is_remote": False,
            "salary_yearly_min": None, "description": ""}
    base.update(kw)
    return base


def test_salary_gate_pass_fail_unknown():
    assert salary_gate(_job(salary_yearly_min=120000), 100000) == "pass"
    assert salary_gate(_job(salary_yearly_min=60000), 100000) == "fail"
    assert salary_gate(_job(), 100000) == "unknown"


def test_low_salary_rejected():
    persona, score, reason, gate = score_job(_job(salary_yearly_min=50000), PROFILE)
    assert gate == "fail" and score == 0.0


def test_title_match_full_containment():
    assert title_match("Senior QA Engineer (Remote)", ["QA Engineer"]) == 1.0


def test_empty_title_scores_zero():
    persona, score, reason, gate = score_job(_job(title=""), PROFILE)
    assert score >= 0  # must not raise


def test_posting_age_days():
    from datetime import datetime, timedelta, timezone
    from pipeline.scoring import posting_age_days
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    old = (datetime.now(timezone.utc) - timedelta(days=20)).strftime("%Y-%m-%d")
    assert posting_age_days(today) == 0
    assert posting_age_days(old) == 20
    assert posting_age_days(None) is None
    assert posting_age_days("") is None
    assert posting_age_days("not-a-date") is None
    assert posting_age_days(today + " 12:00:00") == 0  # timestamp prefix ok


def test_recency_boost_and_stale_penalty():
    from datetime import datetime, timedelta, timezone
    fresh = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    stale = (datetime.now(timezone.utc) - timedelta(days=30)).strftime("%Y-%m-%d")
    _, s_fresh, _, _ = score_job(_job(date_posted=fresh), PROFILE)
    _, s_stale, _, _ = score_job(_job(date_posted=stale), PROFILE)
    assert s_fresh - s_stale == 2.0


def test_tailor_rejects_bad_summary(monkeypatch, tmp_path):
    import pytest
    from pipeline import llm, tailor
    from pipeline.tailor import TailoredResume

    def fake(prompt, schema_model, schema_note, model=None):
        return TailoredResume(summary="too short", experience_bullets=["b"] * 8,
                              skills_lines=["a"] * 4, rationale="r")
    monkeypatch.setattr(llm, "complete_json", fake)
    with pytest.raises(ValueError):
        tailor.tailor_resume({"title": "X", "company": "Y", "description": ""},
                             tmp_path / "out.docx")


def test_tailor_inserts_summary(monkeypatch, tmp_path):
    from docx import Document
    from pipeline import llm, tailor
    from pipeline.tailor import TailoredResume

    summary = ("Quality leader with nearly four years owning end-to-end QA for "
               "healthcare integration software and AI-assisted test automation.")

    def fake(prompt, schema_model, schema_note, model=None):
        return TailoredResume(summary=summary, experience_bullets=["b"] * 8,
                              skills_lines=["a"] * 4, rationale="r")
    monkeypatch.setattr(llm, "complete_json", fake)
    out = tmp_path / "out.docx"
    tailor.tailor_resume({"title": "X", "company": "Y", "description": ""}, out)
    texts = [p.text for p in Document(out).paragraphs]
    assert summary in texts
    assert texts.index(summary) == 2  # right under the contact line


def test_seniority_penalty():
    _, hi, _, _ = score_job(_job(title="QA Engineer"), PROFILE)
    _, lo, _, _ = score_job(_job(title="Principal QA Engineer"), PROFILE)
    assert lo < hi


# ---- packaging helpers ----

def test_slugify_symbols_only():
    assert slugify("!!!***") == "unknown"
    assert slugify("Lower & Co, LLC.") == "lower-co-llc"


def test_clean_letter_strips_wrapper():
    raw = "Here's the cover letter body:\n\n---\n\nDear team, real content.\n\n---\n\nThat's 10 words."
    assert clean_letter(raw) == "Dear team, real content.\n"


def test_clean_letter_passthrough():
    raw = "Dear team, real content.\n"
    assert clean_letter(raw) == raw


def test_clean_letter_single_separator_unchanged():
    raw = "Intro line here that is long enough to keep.\n---\nMore content."
    assert "More content." in clean_letter(raw)


# ---- docx tailoring engine (no LLM - exercises the template surgery) ----

def _template_doc():
    from docx import Document
    from pipeline.tailor import TEMPLATE
    return Document(TEMPLATE)


def test_bullet_blocks_structure():
    from pipeline.tailor import _bullet_blocks
    blocks = _bullet_blocks(_template_doc())
    assert len(blocks) == 3            # experience, skills, leadership
    assert len(blocks[0]) >= 7         # experience bullets
    assert len(blocks[1]) == 4         # skills lines


def test_submit_gate_blocks_unmapped():
    from pipeline.adapters.greenhouse import can_submit
    assert can_submit({"unmapped_required": []}, True) is True
    assert can_submit({"unmapped_required": ["Visa status?*"]}, True) is False
    assert can_submit({"unmapped_required": []}, False) is False
    assert can_submit({}, True) is True


def test_de_ai_strips_dashes():
    from pipeline.destyle import de_ai
    # em dash as sentence separator -> comma
    assert "—" not in de_ai("I own QA—the hard kind—daily.")
    assert de_ai("I own QA—the hard kind—daily.") == "I own QA, the hard kind, daily."
    # numeric range keeps a hyphen, not a comma
    assert de_ai("cut cost 40–50%") == "cut cost 40-50%"
    assert de_ai("cut cost 40—50%") == "cut cost 40-50%"
    # date range
    assert de_ai("2022–Present") == "2022-Present"
    # spaced date range (digit on one side) keeps a hyphen
    assert de_ai("September 2022 – Present") == "September 2022-Present"
    # tight compound proper noun -> hyphen, not comma
    assert de_ai("University of Wisconsin–Madison") == "University of Wisconsin-Madison"
    # en dash and curly quotes
    assert de_ai("Epic’s “Bridges”") == "Epic's \"Bridges\""
    # ellipsis char
    assert de_ai("wait… done") == "wait... done"
    # no dashes at all -> unchanged
    assert de_ai("Plain text, already clean.") == "Plain text, already clean."


def test_de_ai_no_emdash_anywhere():
    from pipeline.destyle import de_ai
    samples = ["Led QA — owned 665 changes", "Python — SQL — Git", "$115,000–$135,000"]
    for s in samples:
        assert "—" not in de_ai(s) and "–" not in de_ai(s)


def test_option_matches_word_boundary():
    from pipeline.adapters.greenhouse import _option_matches
    # "No" must match these
    assert _option_matches("No", "no")
    assert _option_matches("No, I do not", "no")
    assert _option_matches("No.", "no")
    # but NOT these - the bug that selected "Not sure" when we meant "No"
    assert not _option_matches("Not sure", "no")
    assert not _option_matches("None of the above", "no")
    assert not _option_matches("Not applicable", "no")
    # "Yes" boundaries
    assert _option_matches("Yes", "yes")
    assert _option_matches("Yes, authorized", "yes")
    assert not _option_matches("Yesterday", "yes")


def test_submission_confirmed():
    from pipeline.adapters.greenhouse import submission_confirmed
    # positive: a success phrase on the page
    assert submission_confirmed("Thank you for applying to Garner!", False, False)
    assert submission_confirmed("Your application has been received.", False, False)
    # positive: navigated away AND the form is gone
    assert submission_confirmed("", True, True)
    # negative: clicking alone, form still there, no phrase -> NOT confirmed
    assert not submission_confirmed("Please fix the errors below", False, False)
    assert not submission_confirmed("", True, False)   # url changed but form still present
    assert not submission_confirmed("", False, True)   # form gone but no nav/phrase
    assert not submission_confirmed("", False, False)


def test_sponsorship_beats_human_only_guard():
    from pipeline.adapters.greenhouse import is_human_only
    bank = {"requires_sponsorship": "No, I do not now or in the future require sponsorship"}
    # the sponsorship question phrased with 'visa status' must be ANSWERABLE, not blocked
    q = "Will you now or in the future require sponsorship for employment visa status (H-1B)?"
    assert is_human_only(q, "requires_sponsorship", bank) is False
    # but a genuine visa-category question stays human-only
    assert is_human_only("What is your current visa status?", None, bank) is True
    assert is_human_only("Which work authorization best describes you?", "work_authorization", bank) is True
    # EEO stays human-only
    assert is_human_only("What is your gender identity?", None, bank) is True
    # if no sponsorship answer in the bank, don't pretend we can answer
    assert is_human_only(q, "requires_sponsorship", {}) is True


def test_never_autofill_patterns():
    from pipeline.adapters.greenhouse import NEVER_AUTOFILL
    for q in ("Which work authorization best describes you:", "What is your visa status?",
              "Race/Ethnicity", "Veteran status", "Do you have a disability?",
              "Gender identity", "Sexual orientation (Mark all that apply)",
              "I identify as transgender", "Are you LGBTQ+?", "What are your pronouns?",
              "Home Address", "Street Address"):
        assert NEVER_AUTOFILL.search(q), q
    # must NOT block answerable questions
    for ok in ("Are you legally authorized to work in the US?",
               "Will you require sponsorship?", "Email address", "What is your start date?"):
        assert not NEVER_AUTOFILL.search(ok), ok


def test_detect_lane():
    from apply import detect_lane
    assert detect_lane("https://job-boards.greenhouse.io/x/jobs/1") == "auto"
    assert detect_lane("https://www.indeed.com/viewjob?jk=1") == "assisted"
    assert detect_lane(None) == "assisted"


def test_tailor_rejects_malformed_plans(monkeypatch, tmp_path):
    import pytest
    from pipeline import llm, tailor
    from pipeline.tailor import TailoredResume

    def fake(prompt, schema_model, schema_note, model=None):
        return TailoredResume(experience_bullets=["only one"], skills_lines=["a"] * 4,
                              rationale="r")
    monkeypatch.setattr(llm, "complete_json", fake)
    with pytest.raises(ValueError):
        tailor.tailor_resume({"title": "X", "company": "Y", "description": ""},
                             tmp_path / "out.docx")

    def fake2(prompt, schema_model, schema_note, model=None):
        return TailoredResume(experience_bullets=["b"] * 8, skills_lines=["a"] * 2,
                              rationale="r")
    monkeypatch.setattr(llm, "complete_json", fake2)
    with pytest.raises(ValueError):
        tailor.tailor_resume({"title": "X", "company": "Y", "description": ""},
                             tmp_path / "out.docx")


def test_autonomy_gate(tmp_path, monkeypatch):
    import sqlite3
    from pipeline import db as dbmod
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    conn.executescript(dbmod.SCHEMA)
    for m in dbmod.MIGRATIONS:
        try:
            conn.execute(m)
        except sqlite3.OperationalError:
            pass
    assert dbmod.approved_send_count(conn) == 0
    assert dbmod.autonomy_unlocked(conn) is False
    now = "2026-06-12T00:00:00"
    for i in range(dbmod.PROBATION_SENDS):
        conn.execute(
            "INSERT INTO jobs (dedup_hash, title, company, status, send_approved, "
            "first_seen, last_seen) VALUES (?,?,?,?,?,?,?)",
            (f"h{i}", "t", "c", "applied", 1, now, now))
    conn.commit()
    assert dbmod.approved_send_count(conn) == dbmod.PROBATION_SENDS
    assert dbmod.autonomy_unlocked(conn) is True


def test_discover_ats_survives_bad_org():
    from pipeline.discovery import discover_ats
    jobs = discover_ats({"greenhouse_boards": ["this-board-does-not-exist-xyz"],
                         "lever_orgs": [], "ashby_orgs": []})
    assert jobs == []


def test_write_block_shrink_and_grow(tmp_path):
    from pipeline.tailor import _bullet_blocks, _write_block
    doc = _template_doc()
    blocks = _bullet_blocks(doc)
    _write_block(blocks[0], [f"bullet {i}" for i in range(7)])   # shrink 9 -> 7
    _write_block(blocks[1], [f"skill {i}" for i in range(6)])    # grow 4 -> 6
    out = tmp_path / "t.docx"
    doc.save(out)

    from docx import Document
    from pipeline.tailor import _bullet_blocks as blocks_of
    reloaded = blocks_of(Document(out))
    assert [p.text for p in reloaded[0]] == [f"bullet {i}" for i in range(7)]
    assert [p.text for p in reloaded[1]] == [f"skill {i}" for i in range(6)]
    assert len(reloaded[2]) == 2       # leadership untouched
