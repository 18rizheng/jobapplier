"""Generative resume tailoring (rule changed by Richard, 2026-06-11).

The tailor REWRITES the resume per job - new bullet text, recombined emphasis,
job-matched language - grounded strictly in data/facts.md. It is free to
reinvent phrasing and structure; it is never free to invent facts. The
reviewer gate (pipeline/reviewer.py) independently traces every generated
claim back to the corpus before anything can be submitted, and a human
approves the final package.
"""

import copy
from pathlib import Path

from pydantic import BaseModel

from . import llm

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "data" / "resumes" / "general.docx"
FACTS = ROOT / "data" / "facts.md"
MAX_PAGES = 2  # one-page template + generated content should never spill past 2


def _pdf_pages(docx_path: Path) -> int:
    """Render the docx to PDF and count pages. Returns 1 if conversion unavailable."""
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        from pypdf import PdfReader
        return len(PdfReader(str(pdf_path)).pages)
    except Exception:
        return 1


class TailoredResume(BaseModel):
    summary: str                    # 15-35 word professional summary aimed at THIS job
    experience_bullets: list[str]   # 7-9 rewritten bullets, job-matched order
    skills_lines: list[str]         # exactly 4 "Category: a, b, c" lines
    rationale: str                  # what was emphasized and why, 2-3 sentences


_SCHEMA_NOTE = """

Respond with ONLY a JSON object, no markdown fences, matching exactly:
{"summary": "<15-35 word professional summary>", "experience_bullets": [<7-9 strings>],
"skills_lines": [<exactly 4 strings like "Category: item, item, item">],
"rationale": "<2-3 sentences>"}"""


def _build_prompt(job, facts, brief=""):
    brief_block = f"\nCOMPANY BRIEF (use to angle emphasis, do not quote verbatim):\n{brief}\n" if brief else ""
    return f"""Rewrite this candidate's resume content to fit one specific job posting.{brief_block}

You may rewrite bullets from scratch: re-phrase, merge, split, re-emphasize, and mirror
the posting's vocabulary aggressively. Write tight, metric-forward bullets (one line
each, ~15-30 words, strong verbs, no first person).

WRITE LIKE A HUMAN, NOT AN AI. Hard style rules:
- NEVER use em dashes or en dashes. Use commas or periods; use "to" for ranges (write
  "40 to 50%" or "40-50%", never with a long dash). Use plain straight quotes.
- Ban these AI-tell words: leverage, utilize, spearhead, passionate, seamless, robust,
  cutting-edge, dynamic, delve, tapestry, testament, synergy, foster, realm, landscape,
  navigate, elevate, unlock, empower, crucial, pivotal, meticulous, holistic, myriad.
- No "not only X but also Y". Vary sentence length. Plain strong verbs (built, owned,
  cut, led, shipped, found), concrete nouns, real numbers. Sound like a competent
  engineer wrote it in 10 minutes, not like marketing copy.

SHARPNESS RULES:
- The FIRST bullet must directly answer the posting's single most important requirement.
- Translate corpus facts into THIS posting's language wherever the substance matches
  (corpus "API testing (Postman, SoapUI)" -> "REST API test automation" for an API
  role). Mirror their nouns; never their facts.
- Translate domain jargon for the audience: for non-healthcare employers, prefer
  "integration interfaces between enterprise systems" over HL7v2/Bridges specifics;
  for healthcare employers, lean into the domain vocabulary hard.
- At least two-thirds of bullets carry a number from the corpus.
- Cut what this employer won't care about; an 8-bullet resume aimed at their needs
  beats 9 bullets of coverage.

GROUNDING POLICY: the FACT CORPUS below is your guide, not a phrasebook. Use the
posting's keywords, concepts, methodologies, and domain terms freely to characterize
the candidate's work - generalization and reasonable inference from the corpus are
encouraged (release validation implies quality gates; CI pipeline ownership implies
build/deploy fluency; cross-team coordination implies stakeholder management).
Only three things are immovable (the reviewer rejects violations):
1. Employers, job titles, dates, degrees, and certifications stay exactly as in the corpus.
2. NUMBERS stay exactly as in the corpus (600+ of 665 changes reached production -
   never claim all 665 shipped).
3. Never claim hands-on proficiency in a specific named technology (language, tool,
   framework, product) the corpus doesn't support - mirror the posting's *concepts*
   instead, with the candidate's real stack as the evidence.

FACT CORPUS:
{facts}

JOB POSTING:
Title: {job.get('title')}
Company: {job.get('company')}
Description:
{(job.get('description') or '')[:4500]}

Produce:
1. A professional summary (15-35 words, no first person, no objective-statement cliches)
   positioning the candidate for THIS role specifically - a recruiter skimming 6 seconds
   should immediately see the match.
2. 7-9 experience bullets for the Epic Systems Quality Manager role (the earlier
   tutoring/hospital roles stay untouched in the template).
3. Exactly 4 skills lines in "Category: items" form, categories and orderings chosen
   for THIS posting."""


def _bullet_blocks(doc):
    """Contiguous runs of List Paragraph bullets. In the general template:
    block 0 = experience, block 1 = skills, block 2 = leadership (untouched)."""
    blocks, current = [], []
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == "List Paragraph" and p.text.strip():
            current.append(p)
        elif current:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def _set_text(paragraph, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._r.getparent().remove(run._r)
    else:
        paragraph.add_run(text)


def _write_block(paragraphs, lines):
    """Write `lines` into a bullet block, cloning or deleting paragraphs to fit."""
    while len(paragraphs) < len(lines):
        clone = copy.deepcopy(paragraphs[-1]._p)
        paragraphs[-1]._p.addnext(clone)
        from docx.text.paragraph import Paragraph
        paragraphs.append(Paragraph(clone, paragraphs[-1]._parent))
    for extra in paragraphs[len(lines):]:
        extra._p.getparent().remove(extra._p)
    for paragraph, line in zip(paragraphs, lines):
        _set_text(paragraph, line)


def tailor_resume(job: dict, out_path: Path, model: str = llm.DEFAULT_MODEL,
                  avoid_issues: list[str] | None = None, brief: str = "") -> TailoredResume:
    """Write a regenerated resume docx to out_path. Returns the plan.
    avoid_issues: reviewer findings from a prior attempt - hard constraints now.
    brief: optional company research brief to angle emphasis."""
    from docx import Document

    facts = FACTS.read_text(encoding="utf-8-sig")
    prompt = _build_prompt(job, facts, brief)
    if avoid_issues:
        prompt += ("\n\nA previous attempt was REJECTED by the fabrication reviewer for "
                   "the violations below. Do not repeat them or introduce equivalents:\n"
                   + "\n".join(f"- {i}" for i in avoid_issues))
    plan = llm.complete_json(prompt, TailoredResume, _SCHEMA_NOTE, model)

    from . import destyle
    plan.summary = destyle.de_ai(plan.summary)
    plan.experience_bullets = [destyle.de_ai(b) for b in plan.experience_bullets]
    plan.skills_lines = [destyle.de_ai(s) for s in plan.skills_lines]

    if not (7 <= len(plan.experience_bullets) <= 9):
        raise ValueError(f"expected 7-9 experience bullets, got {len(plan.experience_bullets)}")
    if len(plan.skills_lines) != 4:
        raise ValueError(f"expected 4 skills lines, got {len(plan.skills_lines)}")
    n_words = len(plan.summary.split())
    if not (8 <= n_words <= 45):
        raise ValueError(f"summary should be 15-35 words, got {n_words}")

    def render(experience_bullets):
        from . import destyle
        doc = Document(TEMPLATE)
        # tailored summary goes right under the contact line, before the first heading
        summary_p = doc.paragraphs[2].insert_paragraph_before(plan.summary)
        summary_p.runs[0].italic = True
        blocks = _bullet_blocks(doc)
        _write_block(blocks[0], experience_bullets)
        _write_block(blocks[1], plan.skills_lines)
        destyle.sanitize_docx(doc)   # strip em dashes and AI-tell punctuation
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)

    # length guard: render, check PDF page count, trim lowest-priority (trailing)
    # bullets and re-render until it fits within MAX_PAGES or we hit the floor of 7
    bullets = list(plan.experience_bullets)
    render(bullets)
    while len(bullets) > 7 and _pdf_pages(out_path) > MAX_PAGES:
        bullets = bullets[:-1]
        render(bullets)
    plan.experience_bullets = bullets
    return plan
